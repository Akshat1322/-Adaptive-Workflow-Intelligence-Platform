import os
import textwrap
from fpdf import FPDF
from docx import Document

class ReportGenerator:
    """Handles exporting of AI Research Notes to various formats."""
    
    @staticmethod
    def to_pdf(content: str, filename: str) -> str:
        """Export raw markdown-like text to a PDF file."""
        # Clean up some markdown for basic PDF rendering
        lines = content.split('\n')
        
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=11)
        
        def clean_text(text: str) -> str:
            text = text.replace('**', '')
            return text.encode('latin-1', errors='replace').decode('latin-1')

        for line in lines:
            line = clean_text(line)
            if line.startswith('# '):
                pdf.set_font("Arial", 'B', 16)
                pdf.multi_cell(190, 10, txt=line[2:], align='L')
                pdf.set_font("Arial", size=11)
            elif line.startswith('## '):
                pdf.set_font("Arial", 'B', 14)
                pdf.multi_cell(190, 9, txt=line[3:], align='L')
                pdf.set_font("Arial", size=11)
            elif line.startswith('### '):
                pdf.set_font("Arial", 'B', 12)
                pdf.multi_cell(190, 8, txt=line[4:], align='L')
                pdf.set_font("Arial", size=11)
            else:
                for wrapped in textwrap.wrap(line, width=95, break_long_words=True, replace_whitespace=False) or ['']:
                    pdf.multi_cell(190, 7, txt=wrapped)
                
        pdf.output(filename)
        return filename

    @staticmethod
    def to_docx(content: str, filename: str) -> str:
        """Export raw markdown-like text to a DOCX file."""
        doc = Document()
        doc.add_heading('AWIP AI Data Science Report', 0)
        
        lines = content.split('\n')
        for line in lines:
            if line.startswith('# '):
                doc.add_heading(line[2:].replace('**',''), level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:].replace('**',''), level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:].replace('**',''), level=3)
            elif line.startswith('- '):
                doc.add_paragraph(line[2:].replace('**',''), style='List Bullet')
            else:
                if line.strip():
                    doc.add_paragraph(line.replace('**',''))
                    
        doc.save(filename)
        return filename
